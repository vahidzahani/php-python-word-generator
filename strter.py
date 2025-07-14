import os
import json
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image

pic_folder = 'pic'
input_docx = 'asl.docx'
output_docx = 'out.docx'
output_json = 'output.json'

valid_exts = ['.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff']

# Get and sort image files
images = [f for f in os.listdir(pic_folder)
          if os.path.isfile(os.path.join(pic_folder, f)) and os.path.splitext(f)[1].lower() in valid_exts]
images.sort()

# Load base document
doc = Document(input_docx)

cols = 3
rows = (len(images) + cols - 1) // cols
table = doc.add_table(rows=rows, cols=cols)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True
img_width = Inches(1.5)

img_idx = 0
inserted_images = []

for r in range(rows):
    row = table.rows[r]
    for c in range(cols):
        cell = row.cells[c]
        if img_idx < len(images):
            img_path = os.path.join(pic_folder, images[img_idx])
            try:
                with Image.open(img_path) as im:
                    im.verify()
                cell.text = ''
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(img_path, width=img_width)
                inserted_images.append(images[img_idx])
            except Exception:
                cell.text = '[invalid image]'
            img_idx += 1
        else:
            cell.text = ''

doc.save(output_docx)

# Save result as JSON
with open(output_json, 'w', encoding='utf-8') as f:
    json.dump({
        "document": output_docx,
        "images_inserted": inserted_images,
        "total": len(inserted_images)
    }, f, indent=2)

print(f"Document saved as {output_docx}")
print(f"JSON output saved as {output_json}")
